import sys
from docx import Document
from docx.shared import Pt

def convert(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Simple heuristic: keep lines starting and ending with | as table rows
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('|') and '|' in line:
            # gather contiguous table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip().strip('|'))
                i += 1
            # parse header and rows
            if table_lines:
                headers = [h.strip() for h in table_lines[0].split('|')]
                rows = []
                for rl in table_lines[1:]:
                    cells = [c.strip() for c in rl.split('|')]
                    rows.append(cells)
                # create table
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light List Accent 1'
                hdr_cells = table.rows[0].cells
                for idx, h in enumerate(headers):
                    hdr_cells[idx].text = h
                for r in rows:
                    row_cells = table.add_row().cells
                    for idx, c in enumerate(r):
                        if idx < len(row_cells):
                            row_cells[idx].text = c
            continue
        else:
            if line.strip() == '':
                doc.add_paragraph('')
            else:
                p = doc.add_paragraph(line)
            i += 1

    doc.save(output_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: txt_to_docx.py <input.txt> <output.docx>')
        sys.exit(2)
    convert(sys.argv[1], sys.argv[2])
    print('Saved', sys.argv[2])
